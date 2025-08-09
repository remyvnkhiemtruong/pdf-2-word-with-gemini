import sys
import os
import json
import base64
import traceback
import time
import subprocess
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLineEdit, QLabel, QFileDialog, QTextEdit, QMessageBox,
    QListWidget, QTabWidget, QMenuBar
)
from PyQt5.QtCore import QThread, pyqtSignal, QObject, Qt
from PyQt5.QtGui import QColor
from pdf2image import convert_from_path
import google.generativeai as genai
from docx import Document
from PIL import Image

APP_NAME = "GeminiPdfOcrApp"
APP_TITLE = 'PDF to Word by Truong Minh Khiem'
WINDOW_SIZE = (750, 650)
AUTHOR_NAME = "Trương Minh Khiêm"

def get_config_path():
    app_data_path = os.getenv('APPDATA')
    if app_data_path is None:
        app_data_path = os.path.expanduser('~')
    config_dir = os.path.join(app_data_path, APP_NAME)
    os.makedirs(config_dir, exist_ok=True)
    return os.path.join(config_dir, 'config.json')

def get_resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

DARK_STYLE = """
    QWidget { background-color: #2b2b2b; color: #f0f0f0; font-family: Segoe UI; font-size: 10pt; }
    QMainWindow, QMenuBar, QMenuBar::item, QMenu, QMenu::item { background-color: #3c3c3c; color: #f0f0f0; }
    QMenuBar::item:selected, QMenu::item:selected, QTabBar::tab:selected, QTabBar::tab:hover { background: #555555; }
    QTabWidget::pane { border: 1px solid #444444; }
    QTabBar::tab { background: #3c3c3c; color: #f0f0f0; padding: 8px; border: 1px solid #444444; border-bottom: none; }
    QLineEdit, QListWidget { background-color: #3c3c3c; border: 1px solid #555555; }
    QLineEdit { padding: 5px; border-radius: 2px; }
    QTextEdit { background-color: #242424; border: 1px solid #555555; }
    QListWidget { alternate-background-color: #454545; }
    QPushButton { background-color: #555555; border: 1px solid #666666; padding: 8px; border-radius: 4px; }
    QPushButton:hover { background-color: #6a6a6a; }
    QPushButton:pressed { background-color: #4a4a4a; }
    QPushButton#RunButton { background-color: #007acc; font-weight: bold; }
    QPushButton#StopButton { background-color: #d13438; font-weight: bold; }
    QLabel { color: #f0f0f0; }
"""

class OcrWorker(QObject):
    progress = pyqtSignal(str)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)
    file_finished = pyqtSignal(str, bool, str)

    def __init__(self, pdf_paths, api_key, poppler_path, output_dir):
        super().__init__()
        self.pdf_paths = pdf_paths
        self.api_key = api_key
        self.poppler_path = poppler_path
        self.output_dir = output_dir
        self.is_running = True
        self.MAX_RETRIES = 3
        self.RETRY_DELAY = 2

    def run(self):
        try:
            genai.configure(api_key=self.api_key)
            model = genai.GenerativeModel('gemini-2.5-pro')
        except Exception as e:
            self.error.emit(f"Lỗi cấu hình Gemini API: {e}\nKiểm tra lại API Key.")
            return

        total_files = len(self.pdf_paths)
        for file_index, pdf_path in enumerate(self.pdf_paths):
            if not self.is_running:
                self.progress.emit("Đã dừng tác vụ hàng loạt."); return
            
            file_name = os.path.basename(pdf_path)
            self.progress.emit(f"\n--- Bắt đầu xử lý file {file_index + 1}/{total_files}: {file_name} ---")
            
            try:
                images_in_memory = convert_from_path(pdf_path, poppler_path=self.poppler_path)
                num_pages = len(images_in_memory)
                self.progress.emit(f"Chuyển đổi thành công {num_pages} trang.")
                all_markdowns = [self.process_page(model, img, i + 1, num_pages) for i, img in enumerate(images_in_memory) if self.is_running]

                if not self.is_running: continue

                full_markdown = "\n\n---\n\n".join(all_markdowns)
                base_name = os.path.splitext(file_name)[0]
                output_filename = f"{base_name}_ocr.docx"
                output_file_path = os.path.join(self.output_dir, output_filename)
                self.create_word_document(full_markdown, output_file_path)
                self.progress.emit(f"✅ Hoàn thành file: {output_filename}")
                self.file_finished.emit(pdf_path, True, output_file_path)

            except Exception:
                self.progress.emit(f"❌ Lỗi nghiêm trọng khi xử lý file {file_name}: {traceback.format_exc()}")
                self.file_finished.emit(pdf_path, False, "")
                continue
        self.finished.emit("Hoàn thành tất cả các file trong danh sách!")

    def process_page(self, model, pil_image, page_num, total_pages):
        for attempt in range(self.MAX_RETRIES):
            if not self.is_running: return ""
            self.progress.emit(f"Đang xử lý trang {page_num}/{total_pages} (lần thử {attempt + 1})...")
            try:
                prompt = "Hãy thực hiện OCR trên hình ảnh sau và trả về kết quả ở dạng markdown. Giữ nguyên toàn bộ văn bản như trong ảnh, không chỉnh sửa. Nếu có công thức toán học, ghi bằng LaTeX. Nếu có hình ảnh, hãy ghi chú [Hình ảnh] vào đúng vị trí hình đó xuất hiện. Không phân tích lại nội dung."
                response = model.generate_content([prompt, pil_image])
                return response.text
            except Exception as e:
                if attempt >= self.MAX_RETRIES - 1:
                    self.progress.emit(f"Lỗi: Bỏ qua trang {page_num} sau {self.MAX_RETRIES} lần thử thất bại.")
                    return f"\n\n--- Lỗi: Không thể xử lý trang {page_num}, chi tiết: {e} ---\n\n"
                self.progress.emit(f"Cảnh báo (trang {page_num}): Lỗi '{e}'. Thử lại sau {self.RETRY_DELAY} giây.")
                time.sleep(self.RETRY_DELAY)
        return ""

    def create_word_document(self, markdown_text, docx_path):
        doc = Document(); doc.add_heading('Kết quả OCR từ PDF', level=0)
        for line in markdown_text.split('\n'): doc.add_paragraph(line)
        doc.save(docx_path)
    
    def stop(self): self.is_running = False

class App(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(APP_TITLE); self.setGeometry(100, 100, *WINDOW_SIZE)
        self.ocr_thread = None; self.ocr_worker = None
        self.config_path = get_config_path(); self.is_dark_mode = False
        self.setup_ui(); self.load_config()

    def setup_ui(self):
        menu_bar = self.menuBar(); view_menu = menu_bar.addMenu('Giao diện')
        toggle_theme_action = view_menu.addAction('Chuyển chế độ Tối/Sáng')
        toggle_theme_action.triggered.connect(self.toggle_dark_mode)
        self.tabs = QTabWidget(); self.setCentralWidget(self.tabs)
        self.tabs.addTab(self.create_main_tab(), "Xử lý hàng loạt")
        self.tabs.addTab(self.create_settings_tab(), "Cài đặt")
        self.tabs.addTab(self.create_about_tab(), "Thông tin")

    def create_main_tab(self):
        tab = QWidget(); layout = QVBoxLayout(tab)
        layout.addWidget(QLabel("1. Thêm các file PDF cần xử lý:"))
        self.file_list_widget = QListWidget(); self.file_list_widget.setAlternatingRowColors(True)
        layout.addWidget(self.file_list_widget)
        file_buttons_layout = QHBoxLayout()
        add_files_button = QPushButton("Thêm Files PDF..."); add_files_button.clicked.connect(self.add_pdf_files)
        file_buttons_layout.addWidget(add_files_button)
        clear_list_button = QPushButton("Xóa danh sách"); clear_list_button.clicked.connect(self.clear_file_list)
        file_buttons_layout.addWidget(clear_list_button)
        layout.addLayout(file_buttons_layout)
        layout.addWidget(QLabel("\n2. Bắt đầu xử lý:"))
        action_layout = QHBoxLayout()
        self.run_button = QPushButton("🚀 Chạy OCR Hàng Loạt"); self.run_button.setObjectName("RunButton")
        self.run_button.clicked.connect(self.start_ocr)
        action_layout.addWidget(self.run_button)
        self.stop_button = QPushButton("🛑 Dừng"); self.stop_button.setObjectName("StopButton")
        self.stop_button.clicked.connect(self.stop_ocr); self.stop_button.setEnabled(False)
        action_layout.addWidget(self.stop_button)
        layout.addLayout(action_layout)
        layout.addWidget(QLabel("Trạng thái:")); self.status_log = QTextEdit(); self.status_log.setReadOnly(True)
        layout.addWidget(self.status_log)
        return tab

    def create_settings_tab(self):
        tab = QWidget(); layout = QVBoxLayout(tab)
        layout.addWidget(QLabel("<b>Gemini API Key:</b>"))
        self.api_key_edit = QLineEdit(); self.api_key_edit.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.api_key_edit)
        layout.addWidget(QLabel("<b>Thư mục lưu file Word đầu ra:</b>"))
        output_dir_layout = QHBoxLayout()
        self.output_dir_edit = QLineEdit()
        output_dir_layout.addWidget(self.output_dir_edit)
        output_dir_button = QPushButton("Duyệt..."); output_dir_button.clicked.connect(self.select_output_dir)
        output_dir_layout.addWidget(output_dir_button)
        layout.addLayout(output_dir_layout)
        layout.addStretch()
        save_button = QPushButton("Lưu Cài Đặt"); save_button.clicked.connect(self.save_config)
        layout.addWidget(save_button)
        return tab

    def create_about_tab(self):
        tab = QWidget(); layout = QVBoxLayout(tab)
        about_text = f"""<h1 style='text-align:center;'>PDF OCR with Gemini</h1><p style='text-align:center;'>Phiên bản 4.0</p><hr><p><b>Chức năng:</b> Ứng dụng này sử dụng API <b>Google Gemini 1.5 Pro</b> để thực hiện nhận dạng ký tự quang học (OCR) trên các tệp PDF theo hàng loạt.</p><ul><li>Chuyển đổi từng trang PDF thành ảnh.</li><li>Gửi ảnh đến Gemini và nhận về nội dung dạng Markdown.</li><li>Xử lý hàng loạt và tự động lưu kết quả ra file .docx.</li></ul><br><p><b>Tác giả:</b> {AUTHOR_NAME}</p><p><b>Công nghệ sử dụng:</b> Python, PyQt5, Google Generative AI, pdf2image, python-docx.</p>"""
        label = QLabel(about_text); label.setWordWrap(True); label.setAlignment(Qt.AlignTop)
        layout.addWidget(label)
        return tab

    def toggle_dark_mode(self):
        self.is_dark_mode = not self.is_dark_mode
        self.apply_theme()
        self.save_config(show_message=False)

    def apply_theme(self):
        self.setStyleSheet(DARK_STYLE if self.is_dark_mode else "")
    
    def load_config(self):
        if not os.path.exists(self.config_path): return
        try:
            with open(self.config_path, 'r') as f: config = json.load(f)
            api_key_b64 = config.get('api_key', '');
            if api_key_b64: self.api_key_edit.setText(base64.b64decode(api_key_b64).decode('utf-8'))
            self.output_dir_edit.setText(config.get('output_dir', ''))
            self.is_dark_mode = config.get('dark_mode', False)
            self.apply_theme()
        except Exception as e:
            self.status_log.append(f"Lỗi khi tải cấu hình: {e}")
                
    def save_config(self, show_message=True):
        config = {'api_key': base64.b64encode(self.api_key_edit.text().encode('utf-8')).decode('utf-8'), 'output_dir': self.output_dir_edit.text(), 'dark_mode': self.is_dark_mode}
        with open(self.config_path, 'w') as f: json.dump(config, f, indent=4)
        if show_message: QMessageBox.information(self, "Đã lưu", "Đã lưu cài đặt thành công!")

    def add_pdf_files(self):
        fnames, _ = QFileDialog.getOpenFileNames(self, 'Chọn các file PDF', '', 'PDF Files (*.pdf)')
        if fnames:
            for fname in fnames:
                if not self.file_list_widget.findItems(fname, Qt.MatchExactly): self.file_list_widget.addItem(fname)
    
    def clear_file_list(self):
        self.file_list_widget.clear(); self.status_log.clear()

    def select_output_dir(self):
        dname = QFileDialog.getExistingDirectory(self, 'Chọn thư mục lưu file');
        if dname: self.output_dir_edit.setText(dname)

    def start_ocr(self):
        pdf_paths = [self.file_list_widget.item(i).text() for i in range(self.file_list_widget.count())]
        if not pdf_paths:
            QMessageBox.warning(self, "Chưa có file", "Vui lòng thêm ít nhất một file PDF vào danh sách."); return

        api_key = self.api_key_edit.text(); output_dir = self.output_dir_edit.text()
        poppler_path = get_resource_path("poppler_data")

        if not os.path.isdir(poppler_path):
            QMessageBox.critical(self, "Lỗi nghiêm trọng", f"Không tìm thấy thư mục Poppler được đóng gói!\nKiểm tra lại quá trình tạo file.exe.\nĐường dẫn dự kiến: {poppler_path}"); return

        if not all([api_key, output_dir]):
            QMessageBox.warning(self, "Thiếu thông tin", "Vui lòng nhập API Key và chọn Thư mục đầu ra trong tab 'Cài đặt'."); return

        self.run_button.setEnabled(False); self.stop_button.setEnabled(True); self.status_log.clear()
        self.files_to_process_count = len(pdf_paths); self.successful_output_paths = []

        for i in range(self.file_list_widget.count()): self.file_list_widget.item(i).setBackground(QColor(Qt.transparent))

        self.ocr_thread = QThread()
        self.ocr_worker = OcrWorker(pdf_paths, api_key, poppler_path, output_dir)
        self.ocr_worker.moveToThread(self.ocr_thread)
        self.ocr_thread.started.connect(self.ocr_worker.run); self.ocr_worker.finished.connect(self.on_ocr_finished)
        self.ocr_worker.error.connect(self.on_ocr_error); self.ocr_worker.progress.connect(self.status_log.append)
        self.ocr_worker.file_finished.connect(self.on_file_finished)
        self.ocr_thread.start()

    def on_file_finished(self, input_path, success, output_path):
        items = self.file_list_widget.findItems(input_path, Qt.MatchExactly)
        if items:
            item = items[0]
            color = QColor("#2a572a") if success else QColor("#8b0000")
            item.setBackground(color)
            if success and output_path: self.successful_output_paths.append(output_path)

    def stop_ocr(self):
        if self.ocr_worker: self.ocr_worker.stop()
        if self.ocr_thread and self.ocr_thread.isRunning(): self.ocr_thread.quit(); self.ocr_thread.wait()
        self.status_log.append("Đang dừng tác vụ..."); self.run_button.setEnabled(True); self.stop_button.setEnabled(False)

    def on_ocr_finished(self, message):
        if self.files_to_process_count == 1 and self.successful_output_paths: self.open_file(self.successful_output_paths[0])
        QMessageBox.information(self, "Hoàn thành", message); self.cleanup_thread()

    def on_ocr_error(self, error_message):
        QMessageBox.critical(self, "Lỗi", error_message); self.cleanup_thread()
        
    def cleanup_thread(self):
        self.run_button.setEnabled(True); self.stop_button.setEnabled(False)
        self.ocr_thread, self.ocr_worker = None, None

    def open_file(self, path):
        try:
            if sys.platform == "win32": os.startfile(os.path.realpath(path))
            elif sys.platform == "darwin": subprocess.run(["open", path], check=True)
            else: subprocess.run(["xdg-open", path], check=True)
        except Exception as e:
            QMessageBox.warning(self, "Không thể mở file", f"Không thể tự động mở file:\n{path}\n\nLỗi: {e}")

    def closeEvent(self, event):
        if self.ocr_thread and self.ocr_thread.isRunning():
            reply = QMessageBox.question(self, 'Thoát ứng dụng', "Tác vụ đang chạy. Bạn có chắc muốn thoát?", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if reply == QMessageBox.Yes: self.stop_ocr(); event.accept()
            else: event.ignore()
        else: event.accept()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = App()
    ex.show()
    sys.exit(app.exec_())